import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import os

def update_metrics_in_word(excel_path, word_template_path, output_path):
    # Load Excel and normalize column names
    df = pd.read_excel("D://Automation//tests//test.xlsx")
    df = df.dropna(how='all')
    df.columns = df.columns.str.strip().str.lower()

    print("Excel Columns:", df.columns.tolist())

    try:
        app_count = df['application name'].dropna().shape[0]
        total_issues = df['total issues reviewed'].sum()
        actual_issues = df['issue'].sum() - df['probably not an issue'].sum()
        not_an_issue = df['not an issue'].sum()
    except KeyError as e:
        print(f"ERROR: Column not found in Excel: {e}")
        return

    # Map Word metric labels to calculated values
    label_to_logic = {
        "total applications reviewed": app_count,
        "total no.of findings triaged": total_issues,
        "total issues found": actual_issues,
        "total false positive": not_an_issue
    }

    print("Calculated Metrics:", label_to_logic)

    # Load Word document
    doc = Document("D://Automation//tests//metric.docx")

    # Find table with 'Metric' column
    target_table = None
    for table in doc.tables:
        if any(cell.text.strip().lower() == "metric" for cell in table.rows[0].cells):
            target_table = table
            break

    if not target_table:
        print("ERROR: No table with 'Metric' column found.")
        return

    # Identify column indices
    header_cells = target_table.rows[0].cells
    metric_idx = count_idx = None
    for i, cell in enumerate(header_cells):
        col_name = cell.text.strip().lower()
        if col_name == "metric":
            metric_idx = i
        elif col_name == "count":
            count_idx = i

    if metric_idx is None or count_idx is None:
        print("ERROR: Couldn't identify both 'Metric' and 'Count' columns.")
        return

    updated_metrics = {}

    # Update table and capture values for graph
    for row in target_table.rows[1:]:
        metric_text = row.cells[metric_idx].text.strip().lower()
        if metric_text in label_to_logic:
            value = label_to_logic[metric_text]
            row.cells[count_idx].text = str(int(value) if pd.notnull(value) else 0)
            updated_metrics[metric_text] = int(value)
            print(f"✅ Updated '{metric_text}' with value: {value}")
        else:
            print(f"⚠️ Skipped unmatched metric: '{metric_text}'")

    # === Step: Generate a clustered column chart ===
    graph_path = "overview_chart.png"

    bar_labels = [
        "total applications reviewed",
        "total no.of findings triaged",
        "total issues found",
        "total false positive"
    ]
    bar_colors = ["#1f4e79", "#7030a0", "#c00000", "#00b050"]  # dark blue, purple, red, green
    bar_values = [updated_metrics.get(label, 0) for label in bar_labels]

    display_labels = [label.title() for label in bar_labels]

    fig, ax = plt.subplots(figsize=(8, 5))
    x = range(len(display_labels))

    # Plot each bar as a separate group (still one per cluster)
    bars = ax.bar(x, bar_values, color=bar_colors, width=0.6)

    ax.set_xticks(x)
    ax.set_xticklabels(display_labels)
    ax.set_title("Overview")
    ax.set_xlabel("Metric")
    ax.set_ylabel("Count")
    ax.grid(axis='y', linestyle='--', alpha=0.6)

    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, height + 0.5, int(height), ha='center', va='bottom')

    plt.tight_layout()
    plt.savefig(graph_path)
    plt.close()

    # Insert the chart image at the end of the document
    doc.add_paragraph("\n")
    doc.add_paragraph("Overview")
    doc.add_picture(graph_path, width=Inches(5.5))
    print("📊 Graph inserted into the Word document.")

    # Save and clean up
    doc.save("D://Automation//tests//out.docx")
    if os.path.exists(graph_path):
        os.remove(graph_path)

    print(f"\n✅ Word document updated and saved as: {output_path}")


# === Example usage ===
update_metrics_in_word("test.xlsx", "metric.docx", "out1.docx")
